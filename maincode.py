from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, AutoModelForQuestionAnswering,PegasusForConditionalGeneration, PegasusTokenizer,T5ForConditionalGeneration, T5Tokenizer
import fitz
import docx
import streamlit as st
import os
import re
import wikipediaapi
from datasets import load_dataset
from bs4 import BeautifulSoup
import requests
import docxtpl
import io
from docx2pdf import convert


st.set_page_config(
    page_title="QuickRead Q&A",
    page_icon=":telescope:",
    layout="wide",
    initial_sidebar_state="collapsed"
)


# Load the "sciq" dataset
dataset = load_dataset("sciq")

# Define the pipeline for question-answering
qa_tokenizer = AutoTokenizer.from_pretrained("bert-large-uncased-whole-word-masking-finetuned-squad")
qa_model = AutoModelForQuestionAnswering.from_pretrained("bert-large-uncased-whole-word-masking-finetuned-squad")
qa_nlp = pipeline("question-answering", model=qa_model, tokenizer=qa_tokenizer)



# Display selected page
st.sidebar.header("QuickRead")
sidebar_choice = st.sidebar.radio("", ["Home", "About"])

if sidebar_choice == "About":
    st.sidebar.write("App Version:", "QuickRead Q&A v1.0.0")
else:
    st.sidebar.write("Welcome to the QuickRead Q&A app!")     

# Set up main page

st.markdown('<h1 class="title">QuickRead Q&A</h1>', unsafe_allow_html=True)

st.markdown(
    """


## **Interview 1:**
Recently, I had the opportunity of speaking with engineering major and robotics enthusiast Danielle Nillian about the most current developments in robotics technology and how they can affect our daily lives. According to Danielle, robotics technology will fundamentally alter several businesses and every aspect of our life. In this essay, we'll examine some of the recent robotics developments Danielle mentioned as well as some of the moral questions raised by their application.

Danielle cites the development of APIs that make it easier to build and program robots as one of the most significant developments in robotics. Today, anyone can utilize such tools to help them get started with only a little bit of documentation. Now, you can teach your robot to carry out a range of duties, from helping with surgery to delivering packages. Robots will grow more intelligent and autonomous as robotics technology develops, becoming a part of our daily lives and helping us with activities like cooking, cleaning, and even driving.

The rise of collaborative robots, or "cobots," which can collaborate with people in a variety of industries, is another important advancement in robotics. As Danielle pointed out, cobots are made to do dangerous or time-consuming activities, freeing up people to concentrate on more difficult or creative tasks. Cobots are also simple to use and program, making them available to small and medium-sized organizations that would not have the funds to invest in more powerful, expensive robots. Cobots are a game-changer in the industrial sector, improving precision, efficiency, and safety.

But when it comes to using robots, there are some ethical issues to be wary of. As Danielle noted, it is essential to guarantee that robots are created and utilized ethically, taking into account their possible impact on society norms and employment opportunities. Robots run the possibility of being mishandled for immoral activities like monitoring or even violence if they end up in the wrong hands. It is crucial to create rules and regulations for the employment of robots in order to reduce these risks and make sure they uphold societal values and serve the greater good.

In conclusion, robotics technology has made amazing strides in recent years and will have a profound impact on a variety of businesses as well as our daily lives. Robotics developments have the ability to help with daily tasks like cooking and cleaning and can boost production, efficiency, and accuracy in a variety of industries. Cobots can improve security and output in the manufacturing and other sectors of the economy. To make sure that robots adhere to societal ideals and serve society as a whole, ethical considerations must be made when creating, employing, and regulating robots. As Danielle has demonstrated, it's crucial to comprehend both the potential advantages and ethical issues posed by these quickly developing technology.
## **Interview 2:**

Recently, I had the opportunity of speaking with engineering major and robotics enthusiast Danielle Nillian about the most current developments in robotics technology and how they can affect our daily lives. According to Danielle, robotics technology will fundamentally alter several businesses and every aspect of our life. In this essay, we'll examine some of the recent robotics developments Danielle mentioned as well as some of the moral questions raised by their application.

Danielle cites the development of APIs that make it easier to build and program robots as one of the most significant developments in robotics. Today, anyone can utilize such tools to help them get started with only a little bit of documentation. Now, you can teach your robot to carry out a range of duties, from helping with surgery to delivering packages. Robots will grow more intelligent and autonomous as robotics technology develops, becoming a part of our daily lives and helping us with activities like cooking, cleaning, and even driving.

The rise of collaborative robots, or "cobots," which can collaborate with people in a variety of industries, is another important advancement in robotics. As Danielle pointed out, cobots are made to do dangerous or time-consuming activities, freeing up people to concentrate on more difficult or creative tasks. Cobots are also simple to use and program, making them available to small and medium-sized organizations that would not have the funds to invest in more powerful, expensive robots. Cobots are a game-changer in the industrial sector, improving precision, efficiency, and safety.

But when it comes to using robots, there are some ethical issues to be wary of. As Danielle noted, it is essential to guarantee that robots are created and utilized ethically, taking into account their possible impact on society norms and employment opportunities. Robots run the possibility of being mishandled for immoral activities like monitoring or even violence if they end up in the wrong hands. It is crucial to create rules and regulations for the employment of robots in order to reduce these risks and make sure they uphold societal values and serve the greater good.

In conclusion, robotics technology has made amazing strides in recent years and will have a profound impact on a variety of businesses as well as our daily lives. Robotics developments have the ability to help with daily tasks like cooking and cleaning and can boost production, efficiency, and accuracy in a variety of industries. Cobots can improve security and output in the manufacturing and other sectors of the economy. To make sure that robots adhere to societal ideals and serve society as a whole, ethical considerations must be made when creating, employing, and regulating robots. As Danielle has demonstrated, it's crucial to comprehend both the potential advantages and ethical issues posed by these quickly developing technology.

## **What is QuickRead?**

QuickRead is a powerful tool to demonstration how artificial intelligence (AI) can be used to improve productivity in various industries. AI has become increasingly prevalent in recent years, and its potential to enhance productivity is enormous. With the ability to process vast amounts of data quickly and accurately, AI has the potential to streamline workflows, reduce errors, and increase efficiency.

    """,
    unsafe_allow_html=True
)
st.markdown("""
<style>
    .title {
        font-size: 50px;
        color: #2e86c1;
        text-align: center;
        margin-top: 50px;
        margin-bottom: 50px;
    }
    
    .description {
        font-size: 24px;
        color: #333;
        text-align: center;
        margin-bottom: 50px;
    }
    
    .input {
        font-size: 18px;
        color: #333;
        border: 1px solid #ccc;
        border-radius: 5px;
        padding: 10px;
        margin-bottom: 20px;
    }
    
    .button {
        font-size: 18px;
        color: #fff;
        background-color: #2e86c1;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        cursor: pointer;
    }
    
    .button:hover {
        background-color: #1b4f72;
    }
    
    .output {
        font-size: 18px;
        color: #333;
        margin-top: 50px;
        margin-bottom: 50px;
    }
</style>
""", unsafe_allow_html=True)


# Create text inputs
uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])

if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        # Convert PDF to text
        text = ""
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf_reader:
            for page in pdf_reader:
                text += page.get_text()
        context_gpt = text
        
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Convert DOCX to text
        doc = docx.Document(uploaded_file)
        context_gpt = "\n".join([para.text for para in doc.paragraphs])
        
    else:
        st.write("Unsupported file type.")


AI_type = st.selectbox("Use API",["GPT"])

if AI_type == "GPT":
    def generate_text(prompt, context, max_chars=1000000):
        r = requests.post(
            "https://api.deepai.org/api/text-generator",
            data={
                'text': prompt + '\n' + context if context else prompt,
                'max_chars': max_chars
            },
            headers={'api-key': '332c717b-873f-492b-90d5-af813a99ffcf'}
        )
        return r.json()['output']

    prompt = st.text_input("Enter a prompt:")

    # Check if prompt is empty
    if prompt:
        if uploaded_file is not None:
            # Generate text based on prompt and context
            generated_text = generate_text(prompt, context_gpt)

        else:
            # Generate text based on prompt only
            generated_text = generate_text(prompt, "")

        # Display generated text
        st.text_area("Generated Text", value=generated_text, height=400)

        download_docx = st.checkbox("Download as Word document")
        download_pdf = st.checkbox("Download as PDF")

        # Create a new document
        doc = docx.Document()

        # Add a header to the document
        header = doc.sections[0].header
        header.add_paragraph("My Document Header")
        header.add_paragraph("My Name")

         # Add a footer to the document
        footer = doc.sections[0].footer
        footer.add_paragraph("My Document Footer")

        # Insert the generated text into the body of the document
        doc.add_paragraph(generated_text)

        # Save the generated document
        if download_docx:
            # Download button for the generated document
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            st.download_button(
                label="Download Document",
                data=docx_bytes,
                file_name="generated_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        else:
            # Display the generated document
            st.write("Generated Document:")
            st.write(doc)

        # Convert and save the generated document as a PDF
        if download_pdf:
            pdf_bytes = io.BytesIO()
            convert(docx_bytes, pdf_bytes)





context_type = st.selectbox("Choose context type", ["Own text", "Wikipedia","Sciq","Use a Website",""])

if context_type == "Own text":
    # Create text input
    context = st.text_area("Enter your text:")
elif context_type == "Wikipedia":
    # Define Wikipedia API object
    wiki = wikipediaapi.Wikipedia('en')

    # Define list of available contexts
    available_contexts = {
        "Bitcoin": "",
        "Evolution": "",
        "Convolutional neural network": "",
        "World War II":""
    }

    # Extract text from Wikipedia page and update context options
    for context in available_contexts:
        page = wiki.page(context)
        if page.exists():
            available_contexts[context] = page.text

    # Select context from available contexts
    selected_context = st.selectbox("Select a context:", list(available_contexts.keys()))

    # Load selected context into context variable
    context = available_contexts[selected_context]

    # Display context in text area
    st.text_area("Context", value=context, height=200)

if context_type == "Sciq":
    # Select context from sciq dataset
    all_support_text = ' '.join(list(dataset['train']['support']))

    # Load selected context into context variable
    context = all_support_text

    # Display context in text area
    st.text_area("Context", value=context, height=200)

if context_type == "Use a Website":
    # Get user input website URL
    url = st.text_input("Enter website URL:")

    # Check if URL is empty
    if not url:
        st.write("Please enter a website URL.")
    else:
        # Extract text from website
        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        context = soup.get_text()

        # Display context in text area
        st.text_area("Context", value=context, height=200)




# Get user input question
question = st.text_input("Ask your question:")

# Create button
if st.button("Submit"):
    with st.spinner(text="Generating answer..."):
        
        # Generate the answer
        generated = qa_nlp(question=question, context=context, max_answer_len=150)

        # Display the answer
        if generated['score'] > 0.2:
            st.write(f"Answer: {generated['answer']}")
        else:
            st.write("No answer found.")


# Add footer
st.markdown(
    """
    <style>
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #222222;
        color: white;
        text-align: center;
        padding: 10px;
    }
    </style>
    <div class="footer">
        Made with ❤️ by Izzy
    </div>
    """, unsafe_allow_html=True
)